from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "entrega" / "docx"
OUT.mkdir(parents=True, exist_ok=True)

LOGO = ROOT / "docs" / "images" / "logo-icon.png"

PRIMARY = "1F4D78"
ACCENT = "2E74B5"
DARK = "0B2545"
MUTED = "667085"
LIGHT_FILL = "F2F4F7"
CALLOUT = "E8EEF5"
WARN_FILL = "FFF7D6"
BORDER = "D0D5DD"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def wrap_long_value(text: str) -> str:
    if text.startswith("http"):
        return text.replace("/", "/\u200b").replace(".", ".\u200b")
    return text


def set_cell_text(cell, text: str, bold=False, color="000000", size=9.5, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(wrap_long_value(text))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_border_rule(paragraph, color=ACCENT, size="12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def apply_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 28, DARK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 11.5, PRIMARY, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = name.startswith("Heading")


def set_header_footer(doc: Document, label: str) -> None:
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(label)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = "Hermnet | Entrega TFG"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.size = Pt(8)
        fp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document, title: str, subtitle: str, meta: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.55))

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(10)
    kr = kicker.add_run("HERMNET")
    kr.bold = True
    kr.font.size = Pt(11)
    kr.font.color.rgb = RGBColor.from_string(ACCENT)

    t = doc.add_paragraph(style="Title")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run(title)

    st = doc.add_paragraph(style="Subtitle")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(subtitle)

    rule = doc.add_paragraph()
    add_page_border_rule(rule)

    table = doc.add_table(rows=len(meta), cols=2)
    set_table_borders(table, color="E4E7EC", size="4")
    set_table_width(table, [2000, 7000])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_text(row.cells[0], label, bold=True, color=PRIMARY, size=9.5)
        set_cell_shading(row.cells[0], LIGHT_FILL)
        value_size = 8.4 if value.startswith("http") or len(value) > 52 else 9.5
        set_cell_text(row.cells[1], value, size=value_size)

    doc.add_paragraph()
    add_callout(
        doc,
        "Documento preparado para entrega académica. Sustituir los campos [COMPLETAR] por los datos finales antes de exportar a PDF.",
        "Nota de preparación",
        fill=WARN_FILL,
    )


def add_callout(doc: Document, text: str, label: str = "Resumen", fill: str = CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="D0D5DD", size="4")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9000])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(PRIMARY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(DARK)
    return table


def add_manual_toc(doc: Document, headings: list[str]) -> None:
    h = doc.add_paragraph("Índice", style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    for idx, heading in enumerate(headings, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{idx}. {heading}")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(DARK)


def add_section_overview(doc: Document, title: str, items: list[tuple[str, str]]) -> None:
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table, color="E4E7EC", size="4")
    set_table_width(table, [3000, 3000, 3000])
    row = table.rows[0]
    for i, (label, text) in enumerate(items[:3]):
        set_cell_shading(row.cells[i], LIGHT_FILL)
        set_cell_text(row.cells[i], label.upper(), bold=True, color=PRIMARY, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        p = row.cells[i].add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(text)
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor.from_string(DARK)


def read_md(path: str) -> list[str]:
    return (ROOT / path).read_text(encoding="utf-8").splitlines()


def extract_headings(lines: Iterable[str]) -> list[str]:
    out = []
    for line in lines:
        if line.startswith("## "):
            title = re.sub(r"^##\s+\d+\.\s*", "", line).strip()
            if title:
                out.append(title)
    return out


def parse_table(lines: list[str], i: int):
    table_lines = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    if len(table_lines) < 2:
        return None, i
    rows = []
    for idx, line in enumerate(table_lines):
        cells = [c.strip() for c in line.strip("|").split("|")]
        if idx == 1 and all(set(c) <= {"-", ":"} for c in cells):
            continue
        rows.append(cells)
    return rows, i


def add_bullets_as_table_if_reqs(doc: Document, items: list[str]) -> bool:
    if not items:
        return False
    req_like = sum(1 for it in items if re.match(r"R[FN]F\d+", it))
    if req_like < max(2, len(items) // 2):
        return False
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    set_table_width(table, [1700, 7300])
    set_cell_shading(table.rows[0].cells[0], LIGHT_FILL)
    set_cell_shading(table.rows[0].cells[1], LIGHT_FILL)
    set_cell_text(table.rows[0].cells[0], "Código", bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[0].cells[1], "Requisito", bold=True, color=PRIMARY)
    for item in items:
        m = re.match(r"(R[FN]F\d+)\.\s*(.*)", item)
        row = table.add_row()
        if m:
            code, text = m.groups()
        else:
            code, text = "", item
        set_cell_text(row.cells[0], code, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], text, size=9.2)
    return True


def add_markdown_body(doc: Document, md_path: str, skip_title=True) -> None:
    lines = read_md(md_path)
    in_code = False
    code_lines: list[str] = []
    pending_bullets: list[str] = []
    pending_numbers: list[str] = []

    def flush_bullets():
        nonlocal pending_bullets
        if not pending_bullets:
            return
        if not add_bullets_as_table_if_reqs(doc, pending_bullets):
            for item in pending_bullets:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                p.add_run(item)
        pending_bullets = []

    def flush_numbers():
        nonlocal pending_numbers
        for item in pending_numbers:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            p.add_run(item)
        pending_numbers = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_bullets()
                flush_numbers()
                text = "\n".join(code_lines)
                table = doc.add_table(rows=1, cols=1)
                set_table_borders(table, color="D0D5DD", size="4")
                set_table_width(table, [9000])
                cell = table.cell(0, 0)
                set_cell_shading(cell, "F8FAFC")
                set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
                run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                flush_bullets()
                flush_numbers()
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            flush_bullets()
            flush_numbers()
            i += 1
            continue

        if line.startswith("|"):
            flush_bullets()
            flush_numbers()
            rows, ni = parse_table(lines, i)
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=1, cols=cols)
                set_table_borders(table)
                widths = [int(9000 / cols)] * cols
                set_table_width(table, widths)
                for c, text in enumerate(rows[0]):
                    set_cell_shading(table.rows[0].cells[c], LIGHT_FILL)
                    set_cell_text(table.rows[0].cells[c], text, bold=True, color=PRIMARY)
                for r in rows[1:]:
                    row = table.add_row()
                    for c in range(cols):
                        set_cell_text(row.cells[c], r[c] if c < len(r) else "", size=9)
            i = ni
            continue

        if skip_title and line.startswith("# "):
            i += 1
            continue

        if line.startswith("### "):
            flush_bullets()
            flush_numbers()
            doc.add_paragraph(re.sub(r"^###\s+\d+(\.\d+)*\s*", "", line[4:]).strip(), style="Heading 3")
            i += 1
            continue

        if line.startswith("## "):
            flush_bullets()
            flush_numbers()
            doc.add_paragraph(re.sub(r"^##\s+\d+(\.\d+)*\s*", "", line[3:]).strip(), style="Heading 1")
            i += 1
            continue

        if line.startswith("- "):
            flush_numbers()
            pending_bullets.append(line[2:].strip())
            i += 1
            continue

        m_num = re.match(r"^\d+\.\s+(.*)", line)
        if m_num:
            flush_bullets()
            pending_numbers.append(m_num.group(1).strip())
            i += 1
            continue

        flush_bullets()
        flush_numbers()
        p = doc.add_paragraph()
        # Minimal inline bold handling for metadata lines.
        parts = re.split(r"(\*\*.*?\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2])
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(DARK)
            else:
                p.add_run(part)
        i += 1

    flush_bullets()
    flush_numbers()


def build_doc(title: str, subtitle: str, md_path: str, output: str, meta: list[tuple[str, str]], lead: str, overview: list[tuple[str, str]]) -> None:
    doc = Document()
    apply_styles(doc)
    add_cover(doc, title, subtitle, meta)
    headings = extract_headings(read_md(md_path))
    add_manual_toc(doc, headings[:18])
    add_callout(doc, lead, "Resumen ejecutivo")
    add_section_overview(doc, "Mapa del documento", overview)
    add_markdown_body(doc, md_path)
    set_header_footer(doc, title)
    doc.save(OUT / output)


def main() -> None:
    common_meta = [
        ("Proyecto", "Hermnet"),
        ("Repositorio", "https://github.com/Hermnet/Hermnet"),
        ("Alumno/s", "[COMPLETAR]"),
        ("Centro", "[COMPLETAR]"),
        ("Fecha", "[COMPLETAR]"),
    ]
    build_doc(
        "Memoria del Proyecto",
        "Aplicación móvil de mensajería privada con backend propio, cifrado y documentación de entrega",
        "docs/entrega/Memoria_proyecto_Hermnet.md",
        "Memoria_proyecto_Hermnet.docx",
        common_meta + [("Tipo", "Memoria académica")],
        "Documento académico principal del TFG: contextualiza la necesidad, define objetivos, requisitos, planificación, arquitectura, pruebas, sostenibilidad y conclusiones del proyecto Hermnet.",
        [
            ("Enfoque", "Memoria académica y justificación del proyecto"),
            ("Cobertura", "Objetivos, requisitos, planificación, pruebas y conclusiones"),
            ("Entrega", "Preparado para completar datos personales y exportar a PDF"),
        ],
    )
    build_doc(
        "Documentación Técnica",
        "Guía de instalación, arquitectura, backend, frontend, base de datos, pruebas y despliegue",
        "docs/entrega/Documentacion_tecnica_Hermnet.md",
        "Documentacion_tecnica_Hermnet.docx",
        common_meta + [("Tipo", "Documento técnico")],
        "Documento dirigido a evaluadores y desarrolladores: explica cómo instalar, ejecutar, probar, desplegar y mantener el sistema Hermnet desde una perspectiva técnica.",
        [
            ("Backend", "Spring Boot, PostgreSQL, JWT y OpenAPI"),
            ("Frontend", "React Native, Expo, SQLite y servicios móviles"),
            ("Operación", "Instalación, pruebas, errores comunes y mantenimiento"),
        ],
    )
    build_doc(
        "Manual de Usuario",
        "Guía práctica para usar Hermnet: contactos, mensajes, grupos, perfil, seguridad y respaldos",
        "docs/entrega/Manual_usuario_Hermnet.md",
        "Manual_usuario_Hermnet.docx",
        common_meta + [("Tipo", "Manual funcional")],
        "Documento orientado al usuario final: describe de forma ordenada cómo empezar, añadir contactos, enviar mensajes, gestionar grupos, proteger la app y resolver problemas frecuentes.",
        [
            ("Primer uso", "Identidad, PIN y nombre público obligatorio"),
            ("Funciones", "Chats, contactos, grupos, backups y ajustes"),
            ("Soporte", "Problemas frecuentes y buenas prácticas"),
        ],
    )


if __name__ == "__main__":
    main()
